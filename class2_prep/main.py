'''
code programmed by Lupnis <Lupnis@outlook.com>
I write these comments in English simply because Chinese code will turn to
a total mess due to some technical issues.
date: 2021-09-24
**NOTICE** use 'pip install -r requirements.txt' to prepare for the runtime environment
'''
import pdfplumber               # for pdf reading
import re                       # for text matching
import jieba                    # for word-cutting stuff
from docx import *              # for word document saving
from docx.shared import Inches  # for word cloud picture size
from wordcloud import *         # for making wordcloud pictures


def read_pdf(path) -> str:
    txt = ''
    with pdfplumber.open(path) as pdf:
        for i in pdf.pages:
            txt += i.extract_text()
    return txt


def cut_slices(txt) -> str:
    txt = re.findall(r'\w|[\u4E00-\u9FA5]', txt)
    txt = ''.join(txt)
    txt = jieba.lcut(txt)
    ret = ''
    for i in txt:
        ret += str(i)+' '
    return ret


def count_words(txt) -> list:
    dirc = {}
    for t in txt.split(' '):
        if str(t) not in dirc:
            dirc[str(t)] = 1
        else:
            dirc[str(t)] += 1
    dirc = sorted(dirc.items(), key=lambda it: (-it[1], it[0]))
    return dirc


def get_ranks(lst, tms) -> list:
    return lst[0:tms]


def img_generate(content, size, picname):
    word = WordCloud(background_color="black",
                     width=800, height=800,
                     max_words=size,
                     font_path='msyh.ttc').generate(content)
    word.to_file(picname+'.png')


def save_rank(title, lst):
    doc = Document()
    doc.add_heading(title+' statics document', 0)
    doc.add_heading('count list:', 1)
    table = doc.add_table(rows=1, cols=2, style='Light Shading Accent 2')
    table.rows[0].cells[0].text = 'word'
    table.rows[0].cells[1].text = 'count'
    for i in lst:
        single_row = table.add_row().cells
        single_row[0].text = i[0]
        single_row[1].text = str(i[1])

    doc.add_heading('word cloud:', 1)
    doc.add_picture(title+'.png', width=Inches(5))

    doc.save('res_'+title+'.docx')


def main():
    pdf_content = read_pdf('test.pdf')
    content = cut_slices(pdf_content)
    lst = count_words(content)
    lst = get_ranks(lst, 20)
    img_generate(content, 20, 'test')
    save_rank('test', lst)


main()

